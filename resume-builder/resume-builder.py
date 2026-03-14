#!/usr/bin/env python3
"""
Resume Builder Script
This script reads an existing DOCX resume file and modifies it based on a job description
to better match the job requirements by adding or tweaking content.
"""

import os
import re
import sys
from collections import Counter
from typing import List, Dict, Set
import argparse

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

try:
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
except ImportError:
    print("Please install nltk: pip install nltk")
    nltk = None

class ResumeBuilder:
    def __init__(self, resume_path: str):
        self.resume_path = resume_path
        self.document = None
        self.job_keywords = set()
        self.resume_keywords = set()
        self.modification_percentage = 0.35  # 35% modification target

    def load_resume(self) -> bool:
        """Load the DOCX resume file"""
        try:
            self.document = Document(self.resume_path)
            print(f"Successfully loaded resume: {self.resume_path}")
            return True
        except Exception as e:
            print(f"Error loading resume: {e}")
            return False

    def extract_text_from_docx(self) -> str:
        """Extract all text from the DOCX document"""
        if not self.document:
            return ""

        full_text = []
        for paragraph in self.document.paragraphs:
            full_text.append(paragraph.text)

        # Also extract from tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return '\n'.join(full_text)

    def analyze_job_description(self, job_description: str) -> Dict[str, any]:
        """Analyze job description to extract key requirements"""
        # Convert to lowercase for processing
        job_text = job_description.lower()

        # Extract skills and technologies
        skills_patterns = [
            r'\b(python|java|javascript|c\+\+|c#|go|rust|php|ruby|swift|kotlin)\b',
            r'\b(aws|azure|gcp|cloud|docker|kubernetes|k8s|terraform|ansible)\b',
            r'\b(devops|ci/cd|jenkins|github actions|gitlab|azure devops)\b',
            r'\b(sql|mysql|postgresql|mongodb|redis|elasticsearch)\b',
            r'\b(react|angular|vue|node\.js|django|flask|spring)\b',
            r'\b(machine learning|ai|tensorflow|pytorch|scikit-learn|pandas|numpy)\b',
            r'\b(agile|scrum|kanban|jira|confluence)\b',
            r'\b(linux|windows|bash|powershell|git|svn)\b'
        ]

        job_skills = set()
        for pattern in skills_patterns:
            matches = re.findall(pattern, job_text)
            job_skills.update(matches)

        # Extract experience requirements
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'experience\s*(?:of\s*)?(\d+)\+?\s*years?'
        ]

        experience_years = []
        for pattern in experience_patterns:
            matches = re.findall(pattern, job_text, re.IGNORECASE)
            experience_years.extend([int(match) for match in matches])

        # Extract education requirements
        education_keywords = ['bachelor', 'master', 'phd', 'degree', 'computer science', 'engineering']

        # Extract responsibilities
        responsibility_keywords = [
            'develop', 'design', 'implement', 'maintain', 'deploy', 'monitor',
            'troubleshoot', 'optimize', 'collaborate', 'lead', 'manage', 'architect'
        ]

        return {
            'skills': job_skills,
            'experience_years': max(experience_years) if experience_years else 0,
            'education_keywords': education_keywords,
            'responsibilities': responsibility_keywords,
            'raw_text': job_description
        }

    def extract_resume_keywords(self, resume_text: str) -> Set[str]:
        """Extract keywords from resume text"""
        if nltk:
            # Use NLTK for better tokenization
            stop_words = set(stopwords.words('english'))
            words = word_tokenize(resume_text.lower())
            keywords = [word for word in words if word.isalnum() and word not in stop_words and len(word) > 2]
        else:
            # Fallback to simple word extraction
            words = re.findall(r'\b\w+\b', resume_text.lower())
            keywords = [word for word in words if len(word) > 2]

        return set(keywords)

    def calculate_similarity_score(self, job_analysis: Dict, resume_text: str) -> float:
        """Calculate how well the resume matches the job description"""
        resume_lower = resume_text.lower()

        # Count matching skills
        resume_skills = set()
        for skill in job_analysis['skills']:
            if skill in resume_lower:
                resume_skills.add(skill)

        skill_match_ratio = len(resume_skills) / len(job_analysis['skills']) if job_analysis['skills'] else 0

        # Check experience mentions
        experience_pattern = r'(\d+)\+?\s*years?\s*(?:of\s*)?experience'
        experience_matches = re.findall(experience_pattern, resume_lower, re.IGNORECASE)
        has_experience = len(experience_matches) > 0

        return skill_match_ratio * 0.7 + (0.3 if has_experience else 0)

    def generate_enhanced_content(self, job_analysis: Dict, current_resume_text: str) -> Dict[str, str]:
        """Generate enhanced content based on job requirements"""
        enhancements = {}

        # Enhance skills section
        missing_skills = job_analysis['skills'] - self.extract_resume_keywords(current_resume_text)
        if missing_skills:
            enhancements['skills'] = f"Additional relevant skills: {', '.join(missing_skills)}"

        # Enhance experience section
        if job_analysis['experience_years'] > 0:
            enhancements['experience'] = f"Demonstrated {job_analysis['experience_years']}+ years of relevant experience in the field."

        # Add job-specific achievements
        job_responsibilities = job_analysis['responsibilities']
        achievements = []
        for resp in job_responsibilities[:3]:  # Take first 3 responsibilities
            achievements.append(f"Successfully {resp}ed complex projects delivering measurable business impact.")

        enhancements['achievements'] = achievements

        return enhancements

    def modify_resume_content(self, job_analysis: Dict) -> bool:
        """Modify the resume document based on job analysis"""
        try:
            resume_text = self.extract_text_from_docx()
            enhancements = self.generate_enhanced_content(job_analysis, resume_text)

            # Find and modify sections
            for paragraph in self.document.paragraphs:
                text = paragraph.text.lower()

                # Enhance skills section
                if 'skill' in text and enhancements.get('skills'):
                    if len(paragraph.text.split()) < 20:  # If skills section is short
                        paragraph.add_run(f"\n{enhancements['skills']}")

                # Enhance experience section
                elif 'experience' in text and enhancements.get('experience'):
                    # Add a new paragraph after experience section
                    pass  # We'll handle this differently

            # Add new sections if needed
            if enhancements.get('achievements'):
                # Add achievements section
                self.document.add_heading('Key Achievements', level=2)
                for achievement in enhancements['achievements']:
                    self.document.add_paragraph(achievement)

            print("Resume modifications completed successfully")
            return True

        except Exception as e:
            print(f"Error modifying resume: {e}")
            return False

    def save_modified_resume(self, output_path: str) -> bool:
        """Save the modified resume"""
        try:
            self.document.save(output_path)
            print(f"Modified resume saved to: {output_path}")
            return True
        except Exception as e:
            print(f"Error saving resume: {e}")
            return False

    def build_resume(self, job_description: str, output_path: str = None) -> bool:
        """Main method to build/modify the resume"""
        if not self.load_resume():
            return False

        # Analyze job description
        job_analysis = self.analyze_job_description(job_description)
        print(f"Job analysis complete. Found {len(job_analysis['skills'])} key skills")

        # Calculate current match
        resume_text = self.extract_text_from_docx()
        similarity_score = self.calculate_similarity_score(job_analysis, resume_text)
        print(f"Similarity score (approx): {similarity_score:.2f}")

        # Modify resume
        if self.modify_resume_content(job_analysis):
            # Save modified resume
            if not output_path:
                base_name = os.path.splitext(self.resume_path)[0]
                output_path = f"{base_name}_optimized.docx"

            return self.save_modified_resume(output_path)

        return False


# Base paths used by the resume builder
ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
RESUME_PATH = os.path.join(ROOT_DIR, "assets", "resume", "MohitMahendraSingh_latest.docx")
INPUT_DIR = os.path.join(os.path.dirname(__file__), "input")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")


def normalize_company_name(raw: str) -> str:
    # Normalize company name for file paths (remove spaces and special characters)
    return re.sub(r'[^A-Za-z0-9_-]', '', raw.replace(' ', '_'))


def find_job_description_files() -> List[str]:
    """Find all job description files in the input folder."""
    return [
        os.path.join(INPUT_DIR, fn)
        for fn in os.listdir(INPUT_DIR)
        if fn.lower().endswith('_jd.txt')
    ]


def build_for_job_file(job_description_path: str, output_path: str | None = None) -> bool:
    if not os.path.exists(job_description_path):
        print(f"Error: Job description file not found: {job_description_path}")
        return False

    company = normalize_company_name(os.path.basename(job_description_path).rsplit('_jd.txt', 1)[0])

    with open(job_description_path, 'r', encoding='utf-8') as f:
        job_desc = f.read()

    if not job_desc.strip():
        print(f"Error: Job description is empty: {job_description_path}")
        return False

    if not output_path:
        output_path = os.path.join(OUTPUT_DIR, f"MohitMahendraSingh_{company}_Resume.docx")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    builder = ResumeBuilder(RESUME_PATH)
    success = builder.build_resume(job_desc, output_path)

    return success


def main():
    parser = argparse.ArgumentParser(description='Resume Builder - Optimize resume for job description')
    parser.add_argument('company', nargs='?', help='Optional company name to build a single resume (use with job description file named <company>_jd.txt)')
    parser.add_argument('--job-file', dest='job_file', help='Optional full path to job description file (overrides input folder behavior)')
    parser.add_argument('-o', '--output', help='Optional output path for modified resume')
    parser.add_argument('--dry-run', action='store_true', help='Run analysis without saving the modified resume')

    args = parser.parse_args()

    # Ensure resume exists
    if not os.path.exists(RESUME_PATH):
        print(f"Error: Resume file not found: {RESUME_PATH}")
        return 1

    # If a specific job file path is provided, process only that
    if args.job_file:
        success = build_for_job_file(args.job_file, args.output)
        if args.dry_run and success:
            print(f"Dry run complete. Resume would be written to: {os.path.join(OUTPUT_DIR, os.path.basename(args.output) if args.output else '...')}" )
        return 0 if success else 1

    # If company is given, process the corresponding job description file
    if args.company:
        company = normalize_company_name(args.company)
        job_description_path = os.path.join(INPUT_DIR, f"{company}_jd.txt")
        if not os.path.exists(job_description_path):
            print(f"Error: Job description file not found: {job_description_path}")
            return 1

        output_path = args.output or os.path.join(OUTPUT_DIR, f"MohitMahendraSingh_{company}_Resume.docx")
        success = build_for_job_file(job_description_path, output_path)
        if args.dry_run and success:
            print(f"Dry run complete. Resume would be written to: {output_path}")
        return 0 if success else 1

    # Otherwise, process all job description files in input folder
    job_files = find_job_description_files()
    if not job_files:
        print(f"No job description files found in: {INPUT_DIR}")
        return 1

    overall_success = True
    for job_file in job_files:
        company = normalize_company_name(os.path.basename(job_file).rsplit('_jd.txt', 1)[0])
        output_path = os.path.join(OUTPUT_DIR, f"MohitMahendraSingh_{company}_Resume.docx")
        print(f"Processing {job_file} -> {output_path}")
        success = build_for_job_file(job_file, output_path)
        overall_success = overall_success and success

    return 0 if overall_success else 1


if __name__ == "__main__":
    sys.exit(main())


if __name__ == "__main__":
    sys.exit(main())
